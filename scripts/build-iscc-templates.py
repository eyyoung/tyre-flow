#!/usr/bin/env python3
"""Regenerate template/ISCC_PLUS.docx (ISCC PLUS v2.0) and template/ISCC_EU.docx
(ISCC EU v2.3) from the layout of the official ISCC PDF forms in template/.

The generated files are docxtemplater templates used by src/lib/iscc-export-generator.ts.
They can also be edited directly in Word; this script only exists to rebuild them
from scratch in a reproducible way.

Usage:
    python3 -m venv .venv && .venv/bin/pip install python-docx pypdf
    .venv/bin/python scripts/build-iscc-templates.py            # writes to template/
    .venv/bin/python scripts/build-iscc-templates.py /some/dir  # writes elsewhere

Placeholders (docxtemplater syntax):
  {storeName} {address} {postcodeCity} {cityPostcode} {country} {phone}
  {geoCoordinates} {collectionPoint} {legalPerson} {position} {placeDate}
  {minVolumeCheck} {maxCapacity} {maxSustainableCapacity} (EU only)
  {%signature} (docxtemplater image module)
"""
import io
import os
import sys
import zipfile
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Mm, RGBColor

REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE_DIR = os.path.join(REPO_ROOT, "template")
OUT_DIR = sys.argv[1] if len(sys.argv) > 1 else TEMPLATE_DIR


def extract_logo():
    """Pull the ISCC logo out of the official PDF so no extra asset is needed."""
    from pypdf import PdfReader
    reader = PdfReader(os.path.join(TEMPLATE_DIR, "ISCC_PLUS.pdf"))
    image = reader.pages[0].images[0]
    return io.BytesIO(image.data)


LOGO = extract_logo()
BODY_FONT = "Arial"
EAST_ASIA_FONT = "Microsoft YaHei"
GREY = "5F5F5F"
# Checkboxes are Wingdings glyphs stored the way Word stores symbol-font
# characters (private-use codes U+F0xx). Wingdings is declared in fontTable.xml
# with the symbol charset so that LibreOffice (which has no Wingdings) recodes
# the glyphs to its bundled OpenSymbol font instead of printing letters.
CHECKBOX_FONT = "Wingdings"
BOX_EMPTY = "\uf0a8"    # ☐ (Wingdings 0xA8)
BOX_CHECKED = "\uf0fe"  # ☑ (Wingdings 0xFE)
WINGDINGS_FONT_DECL = (
    '<w:font w:name="Wingdings"><w:panose1 w:val="05000000000000000000"/>'
    '<w:charset w:val="02"/><w:family w:val="auto"/><w:pitch w:val="variable"/>'
    '<w:sig w:usb0="00000000" w:usb1="10000000" w:usb2="00000000" w:usb3="00000000" '
    'w:csb0="80000000" w:csb1="00000000"/></w:font>'
)


# ----------------------------------------------------------------- helpers
def set_run_fonts(run, name=BODY_FONT, east_asia=EAST_ASIA_FONT):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)
    rfonts.set(qn("w:eastAsia"), east_asia)


def add_run(p, text, size=10, bold=False, italic=False, color=None, underline=False,
            superscript=False, font=BODY_FONT, east_asia=EAST_ASIA_FONT):
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if superscript:
        run.font.superscript = True
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    set_run_fonts(run, font, east_asia)
    return run


def add_checkbox(p, checked=False, size=10):
    """Checkbox glyph (Wingdings empty box / box with X)."""
    return add_run(p, BOX_CHECKED if checked else BOX_EMPTY, size=size,
                   font=CHECKBOX_FONT, east_asia=CHECKBOX_FONT)


def declare_symbol_font(path):
    """Append the Wingdings symbol-font declaration to word/fontTable.xml."""
    tmp = path + ".tmp"
    with zipfile.ZipFile(path) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/fontTable.xml":
                text = data.decode("utf8")
                if 'w:name="Wingdings"' not in text:
                    text = text.replace("</w:fonts>", WINGDINGS_FONT_DECL + "</w:fonts>")
                data = text.encode("utf8")
            zout.writestr(item, data)
    os.replace(tmp, path)


def para_format(p, before=0, after=0, line=None, align=None, left=None, first=None,
                keep_next=False, keep_together=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line
    if align is not None:
        p.alignment = align
    if left is not None:
        pf.left_indent = Pt(left)
    if first is not None:
        pf.first_line_indent = Pt(first)
    pf.keep_with_next = keep_next
    pf.keep_together = keep_together
    return p


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcpr.append(shd)


def set_cell_valign(cell, val="center"):
    tcpr = cell._tc.get_or_add_tcPr()
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), val)
    tcpr.append(v)


def set_cell_borders(cell, **edges):
    """edges: top/bottom/left/right -> dict(val, sz, color) or None (nil)."""
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcpr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        if edge not in edges:
            continue
        spec = edges[edge]
        el = OxmlElement(f"w:{edge}")
        if spec is None:
            el.set(qn("w:val"), "nil")
        else:
            el.set(qn("w:val"), spec.get("val", "single"))
            el.set(qn("w:sz"), str(spec.get("sz", 4)))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), spec.get("color", "000000"))
        borders.append(el)


def set_table_borders(table, sz=4, color="000000", val="single"):
    tblpr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblpr.append(borders)


def set_table_cell_margins(table, top=40, bottom=40, left=90, right=90):
    tblpr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for edge, w in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(w))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblpr.append(mar)


def set_table_fixed(table, widths_pt):
    """Fixed layout with explicit grid + cell widths (points)."""
    table.autofit = False
    tblpr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblpr.append(layout)
    # total width
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(int(sum(widths_pt) * 20)))
    tblw.set(qn("w:type"), "dxa")
    # grid
    grid = table._tbl.tblGrid
    for gc in list(grid):
        grid.remove(gc)
    for w in widths_pt:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w * 20)))
        grid.append(gc)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths_pt[idx])


def set_row_height(row, pt, rule=WD_ROW_HEIGHT_RULE.AT_LEAST):
    row.height = Pt(pt)
    row.height_rule = rule


def set_row_cant_split(row):
    trpr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:cantSplit")
    trpr.append(el)


def merge_row(table, r, c_from, c_to):
    a = table.cell(r, c_from)
    b = table.cell(r, c_to)
    return a.merge(b)


def new_document(margins_mm, footer_mm):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Mm(margins_mm["top"])
    sec.bottom_margin = Mm(margins_mm["bottom"])
    sec.left_margin = Mm(margins_mm["left"])
    sec.right_margin = Mm(margins_mm["right"])
    sec.header_distance = Mm(10)
    sec.footer_distance = Mm(footer_mm)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), BODY_FONT)
    rfonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    # python-docx's default template uses theme fonts; drop them so Arial wins.
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        if rfonts.get(qn(attr)) is not None:
            del rfonts.attrib[qn(attr)]
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0
    return doc


def add_field(p, instruction, placeholder, size=10):
    """Insert a complex field (PAGE, NUMPAGES, ...) with a cached placeholder."""
    def fld_char(kind):
        r = add_run(p, "", size=size)
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), kind)
        r._element.append(fc)
    fld_char("begin")
    r = add_run(p, "", size=size)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    r._element.append(instr)
    fld_char("separate")
    add_run(p, placeholder, size=size)
    fld_char("end")


def restart_page_numbers(doc):
    """<w:pgNumType w:start="1"/> so every declaration counts from page 1
    when the export merges several of them into one document."""
    sect_pr = doc.sections[0]._sectPr
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:start"), "1")
    sect_pr.append(pg)


def build_footer(doc, version_text, text_width_pt, logo_height_mm=12.3,
                 page_numbers=False):
    """Footer as in the official forms: ISCC logo bottom-left with the copyright
    line under it, version text right-aligned on the logo's baseline and, for
    the EU form, a centred "1 of 2" page counter."""
    sec = doc.sections[0]
    ftr = sec.footer
    ftr.is_linked_to_previous = False
    # Borderless table: right-aligned tab stops at the margin are not
    # honoured reliably by LibreOffice, a table is.
    cols = 3 if page_numbers else 2
    t = ftr.add_table(rows=2, cols=cols, width=Pt(text_width_pt))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(t, val="nil")
    set_table_cell_margins(t, top=0, bottom=0, left=0, right=0)
    if page_numbers:
        set_table_fixed(t, [text_width_pt * 0.4, text_width_pt * 0.2, text_width_pt * 0.4])
        logo_cell, page_cell, version_cell = t.rows[0].cells
        set_cell_valign(page_cell, "bottom")
        pp = page_cell.paragraphs[0]
        para_format(pp, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
        # "1 of 2": page number restarts per section (see restart_page_numbers)
        # while the total is fixed text, so the batch merge only has to insert
        # section breaks between declarations.
        add_field(pp, "PAGE", "1", size=11)
        add_run(pp, " of 2", size=11)
    else:
        set_table_fixed(t, [text_width_pt / 2, text_width_pt / 2])
        logo_cell, version_cell = t.rows[0].cells
    set_cell_valign(logo_cell, "bottom")
    p1 = logo_cell.paragraphs[0]
    para_format(p1, after=0)
    LOGO.seek(0)
    p1.add_run().add_picture(LOGO, height=Mm(logo_height_mm))
    set_cell_valign(version_cell, "bottom")
    p2 = version_cell.paragraphs[0]
    para_format(p2, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_run(p2, version_text, size=8)
    copyright_cell = t.rows[1].cells[0]
    p3 = copyright_cell.paragraphs[0]
    para_format(p3, before=2, after=0, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_run(p3, "© ISCC System GmbH", size=8)
    # python-docx leaves an empty paragraph before the table; drop it and end
    # the footer with a minimal paragraph instead (Word expects one after a table).
    first = ftr.paragraphs[0]
    if first._p is t._tbl.getprevious() and not first.text:
        first._p.getparent().remove(first._p)
    tail = ftr.add_paragraph()
    para_format(tail, after=0)
    tail_rpr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "2")
    tail_rpr.append(sz)
    tail._p.get_or_add_pPr().append(tail_rpr)


def add_title_table(doc, left_text, right_text, widths_pt):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(t, sz=8, color=GREY)
    set_table_cell_margins(t, top=60, bottom=60, left=100, right=100)
    set_table_fixed(t, widths_pt)
    set_row_height(t.rows[0], 22)
    for cell, text in zip(t.rows[0].cells, (left_text, right_text)):
        set_cell_shading(cell, GREY)
        set_cell_valign(cell, "center")
        p = cell.paragraphs[0]
        para_format(p, after=0)
        add_run(p, text, bold=True, color="FFFFFF")
    return t


def label_cell(cell, text, bold=False, extra_lines=(), valign="center"):
    set_cell_valign(cell, valign)
    p = cell.paragraphs[0]
    para_format(p, after=0)
    add_run(p, text, bold=bold)
    for line in extra_lines:
        p2 = cell.add_paragraph()
        para_format(p2, after=0)
        add_run(p2, line)
    return cell


def placeholder_cell(cell, tag, valign="center"):
    set_cell_valign(cell, valign)
    p = cell.paragraphs[0]
    para_format(p, after=0)
    add_run(p, tag)
    return cell


def checkbox_lines(cell, items, valign="center"):
    """items: list of (checked, label)."""
    set_cell_valign(cell, valign)
    first = True
    for checked, label in items:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        para_format(p, after=0)
        add_checkbox(p, checked)
        add_run(p, " " + label)
    return cell


def add_numbered_item(doc, number, text_runs, num_indent, text_indent, after=6,
                      keep_next=False):
    """text_runs: list of (text, bold) tuples. Hanging-indent numbered paragraph."""
    p = doc.add_paragraph()
    para_format(p, after=after, left=text_indent, first=num_indent - text_indent,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY, keep_next=keep_next)
    p.paragraph_format.tab_stops.add_tab_stop(Pt(text_indent), WD_TAB_ALIGNMENT.LEFT)
    add_run(p, f"{number}.\t")
    for text, bold in text_runs:
        add_run(p, text, bold=bold)
    return p


def add_sub_paragraph(doc, text_runs, text_indent, after=6, keep_next=False):
    p = doc.add_paragraph()
    para_format(p, after=after, left=text_indent, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                keep_next=keep_next)
    for text, bold in text_runs:
        add_run(p, text, bold=bold)
    return p


def add_signature_block(doc, widths_pt, name_function_tag):
    """Place, Date / Signature line + Full name and function line.

    Borderless table so the signature image sits in its own cell.
    """
    t = doc.add_table(rows=2, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(t, val="nil")
    set_table_cell_margins(t, top=20, bottom=20, left=60, right=60)
    set_table_fixed(t, widths_pt)
    set_row_height(t.rows[0], 34)
    set_row_cant_split(t.rows[0])
    set_row_cant_split(t.rows[1])

    line = {"val": "single", "sz": 6, "color": "000000"}
    r0 = t.rows[0].cells
    label_cell(r0[0], "Place, Date:", bold=True, valign="bottom")
    placeholder_cell(r0[1], "{placeDate}", valign="bottom")
    set_cell_borders(r0[1], bottom=line)
    label_cell(r0[2], "Signature:", bold=True, valign="bottom")
    sig = r0[3]
    set_cell_valign(sig, "bottom")
    sp = sig.paragraphs[0]
    para_format(sp, after=0)
    add_run(sp, "{%signature}")
    set_cell_borders(sig, bottom=line)

    r1 = t.rows[1].cells
    c = merge_row(t, 1, 0, 1)
    set_cell_valign(c, "bottom")
    p = c.paragraphs[0]
    para_format(p, before=10, after=0)
    add_run(p, "Full name and function of signatory:", bold=True)
    c2 = merge_row(t, 1, 2, 3)
    set_cell_valign(c2, "bottom")
    p2 = c2.paragraphs[0]
    para_format(p2, before=10, after=0)
    add_run(p2, name_function_tag)
    set_cell_borders(c2, bottom=line)
    # keep the whole block on one page
    for row in t.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.keep_with_next = True
                para.paragraph_format.keep_together = True
    return t


# ---------------------------------------------------------- shared texts
INFO_HEADER = ("Information about the Point of Origin (hereinafter referred to as "
               "“Point of Origin” or “We”):")
GEO_LABEL = "Geo-coordinates (Latitude, Longitude)"
GEO_HINT = ["Required format: must be in decimal degrees, for",
            "example: (-)XX.YYYYYY, (-)XX.YYYYYY"]
MATERIAL_LABEL = "The delivered material consists of the following waste or residues:"
MATERIAL_TEXT = "Biogenic fraction of end-of-life tires"
MATERIAL_NOTE = ("Note: List each waste or residue delivered. Identify each clearly, and give "
                 "the waste codes (if applicable) according to the relevant national waste "
                 "ordinance - if you are entitled to do so (and if applicable).")

ITEM_1 = ("We confirm compliance with all legal obligations as well as the relevant ISCC "
          "(ISCC System GmbH) requirements*. (e.g. for quantities delivered under ISCC) "
          "including contractual agreements with subcontractors and recipients (Collecting "
          "Points), delivery notes/ weighbridge tickets.")
ITEM_2 = ("The material supplied under this self-declaration meets the definition of "
          "“waste” or “residue”.")
ITEM_9 = ("If audits of Certification Bodies or ISCC reveal that relevant ISCC requirements "
          "are not complied with or declarations made in this self-declaration are not "
          "correct, and if the Point of Origin is thereupon excluded as supplier of ISCC "
          "certified material, ISCC is entitled to publish the exclusion of the Point of "
          "Origin on the ISCC website.")
ITEM_ACK = ("We acknowledge and agree that any information relating to Us that We disclose "
            "to other ISCC-certified elements of the supply chain may be further disclosed "
            "by those elements of the supply chain to their Certification Bodies and to ISCC.")
ITEM_WARRANT = ("We warrant that We have a valid legal basis, or have obtained consent from "
                "the natural persons whose personal data (e.g. name, contact details) is "
                "included in this self-declaration, to include such personal data herein and "
                "to disclose and forward it in accordance with the terms set out in this "
                "self-declaration.")
ITEM_PROVIDE = ("We will provide any documentation reasonably required to support the "
                "information contained in this self-declaration to any relevant element of "
                "the supply chain, the Certification Body, ISCC, or any competent authority "
                "or supervisory body immediately upon request. This obligation continues for "
                "five (5) years after the expiry of this self-declaration.")
ITEM_CORRECT = ("All information contained in this self-declaration is correct, up to date, "
                "complete, fully documented, and a fair representation of actual facts. Such "
                "documentation must be kept available for five (5) years after the expiry of "
                "this self-declaration.")
ITEM_LAW = ("This self-declaration and any dispute relating to declarations or information "
            "contained in this self-declaration and its use shall be exclusively be governed "
            "by and construed in accordance with the laws of the Federal Republic of Germany "
            "without giving effect to any conflicts of law principles or rules, and excluding "
            "the application of the United Nations Convention on Contracts for the "
            "International Sale of Goods (CISG). The competent courts in Cologne, Germany, "
            "shall have exclusive jurisdiction for any dispute relating to declarations or "
            "information contained in this self-declaration and its use.")
FOOTNOTE_STAR = ("The ISCC requirements and system documents are available on the ISCC "
                 "website (www.iscc-system.org).")


def add_signing_sentence(doc, after=8):
    p = doc.add_paragraph()
    para_format(p, before=10, after=after, align=WD_ALIGN_PARAGRAPH.JUSTIFY, keep_next=True)
    add_run(p, "By signing this self-declaration, I, ", bold=True)
    add_run(p, "{legalPerson}", bold=True, underline=True)
    add_run(p, ", acting in my capacity as ", bold=True)
    add_run(p, "{position}", bold=True, underline=True)
    add_run(p, " and authorised representative of the Point of Origin, hereby declare, "
               "confirm and agree to the following on behalf of the Point of Origin:",
            bold=True)
    return p


# ================================================================ ISCC PLUS
def build_plus(path):
    # PDF: left 36.2pt, right 595-562=33pt, title top 79.6pt,
    # footer logo 767..802pt from top, copyright line at 810pt
    text_w = 525.8  # 562 - 36.2
    doc = new_document({"top": 28, "bottom": 30, "left": 12.7, "right": 11.7},
                       footer_mm=8.5)
    build_footer(doc, "Version 2.0, as of 30 September 2025", text_w)

    add_title_table(doc, "ISCC PLUS self-declaration",
                    "Point of Origin for waste and residues",
                    [213.1, text_w - 213.1])
    para_format(doc.add_paragraph(), after=0)  # spacer (PDF gap ~14pt)

    # grid columns from PDF x: 36.2 | 301 | 419.1 | 485.7 | 521.6 | 562
    widths = [264.8, 118.1, 66.6, 35.9, 40.4]
    t = doc.add_table(rows=12, cols=5)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(t, sz=4)
    set_table_cell_margins(t, top=50, bottom=50, left=100, right=100)
    set_table_fixed(t, widths)

    # R0 header
    c = merge_row(t, 0, 0, 4)
    label_cell(c, INFO_HEADER, bold=True)
    set_row_height(t.rows[0], 18)

    simple_rows = [
        (1, "Site name", "{storeName}"),
        (2, "Street address", "{address}"),
        (3, "Postcode, City, Country", "{postcodeCity}, {country}"),
        (4, "Phone number", "{phone}"),
    ]
    for r, label, tag in simple_rows:
        label_cell(t.cell(r, 0), label)
        placeholder_cell(merge_row(t, r, 1, 4), tag)
        set_row_height(t.rows[r], 22.8)

    # R5 geo
    label_cell(t.cell(5, 0), GEO_LABEL, extra_lines=GEO_HINT)
    placeholder_cell(merge_row(t, 5, 1, 4), "{geoCoordinates}")
    set_row_height(t.rows[5], 43)

    # R6 material (merged)
    c = merge_row(t, 6, 0, 4)
    set_cell_valign(c, "top")
    p = c.paragraphs[0]
    para_format(p, after=4)
    add_run(p, MATERIAL_LABEL)
    p = c.add_paragraph()
    para_format(p, before=10, after=12, left=6)
    add_run(p, MATERIAL_TEXT)
    p = c.add_paragraph()
    para_format(p, before=2, after=0)
    add_run(p, MATERIAL_NOTE)
    set_row_height(t.rows[6], 70)

    # R7 volume statement + checkbox
    c = merge_row(t, 7, 0, 3)
    label_cell(c, "The Point of Origin generates at least ten (10) metric tons of waste per "
                  "month, or at least one hundred twenty (120) metric tons per year. This "
                  "includes all waste volumes, indented or not for ISCC certification.")
    cb = t.cell(7, 4)
    set_cell_valign(cb, "center")
    p = cb.paragraphs[0]
    para_format(p, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "{minVolumeCheck}", font=CHECKBOX_FONT, east_asia=CHECKBOX_FONT)
    set_row_height(t.rows[7], 52)

    # R8 recipient
    label_cell(t.cell(8, 0), "Recipient of the waste or residue (Collecting Point):")
    placeholder_cell(merge_row(t, 8, 1, 4), "{collectionPoint}")
    set_row_height(t.rows[8], 22.8)

    # R9 plastic RIC
    c = merge_row(t, 9, 0, 1)
    label_cell(c, "In case of plastic waste: The delivery is essentially free of paper, "
                  "biomass and/ or used tires and consists of plastic of the Resin "
                  "Identification Code (RIC) categories (please tick boxes):")
    checkbox_lines(t.cell(9, 2), [(False, "1 PET"), (False, "3 PVC"),
                                  (False, "5 PP"), (False, "7 Other")])
    checkbox_lines(merge_row(t, 9, 3, 4), [(False, "2 HDPE"), (False, "4 LDPE"),
                                           (False, "6 PS")])
    set_row_height(t.rows[9], 59)

    # R10 type of waste
    c = merge_row(t, 10, 0, 1)
    label_cell(c, "Type of waste/residues material")
    checkbox_lines(merge_row(t, 10, 2, 4), [(True, "Post-consumer material"),
                                            (False, "Pre-consumer material"),
                                            (False, "Mixed/unspecified")])
    set_row_height(t.rows[10], 40)

    # R11 classification
    c = merge_row(t, 11, 0, 1)
    set_cell_valign(c, "center")
    p = c.paragraphs[0]
    para_format(p, after=0)
    add_run(p, "ISCC PLUS raw material classification of the waste or residue")
    add_run(p, "1", superscript=True)
    checkbox_lines(merge_row(t, 11, 2, 4), [(False, "Circular"), (True, "Bio-circular")])
    set_row_height(t.rows[11], 33)

    for row in t.rows:
        set_row_cant_split(row)

    # signing sentence + numbered declarations
    add_signing_sentence(doc)
    NUM, TXT = 4, 22  # PDF: number at x=40, text at x=58 (margin 36)
    add_numbered_item(doc, 1, [(ITEM_1, False)], NUM, TXT)
    add_numbered_item(doc, 2, [(ITEM_2, False)], NUM, TXT, after=3, keep_next=True)
    add_sub_paragraph(doc, [("A ", False), ("waste", True),
                            (" means any substance or object which the holder discards, "
                             "intends or is required to discard. This material has reached "
                             "the end of its intended life cycle. The waste was not "
                             "intentionally produced, and its further use requires an "
                             "additional processing step.", False)], TXT, after=3)
    add_sub_paragraph(doc, [("A ", False), ("residue", True),
                            (" means a substance that is not the end product(s) that a "
                             "production process directly seeks to produce; it is not a "
                             "primary aim of the production process, and the process has not "
                             "been deliberately modified to produce it.", False)], TXT)
    add_numbered_item(doc, 3, [("The material supplied under this self-declaration complies "
                                "with the requirements (if any) set on the ISCC PLUS Material "
                                "List (e.g. for UCO, Food Waste, or Silicon Waste).", False)],
                      NUM, TXT)
    add_numbered_item(doc, 4, [("The Point of Origin holds appropriate licenses and permits to "
                                "act as a legal waste management company or is an entity that "
                                "generates recovered material as defined in ISO 14021:2021. In "
                                "either case, this can be proven by relevant documentation. "
                                "Recovered material is defined by ISO as material that would "
                                "have otherwise been disposed of as waste or used for energy "
                                "recovery but has instead been collected and recovered as a "
                                "material input instead of using new primary material for a "
                                "recycling or manufacturing process.", False)], NUM, TXT)
    add_numbered_item(doc, 5, [("Compliance with applicable national and regional legislation "
                                "is ensured (in particular with respect to the definition of "
                                "waste, waste prevention, waste collection, waste sorting, "
                                "transport, labelling of waste, etc.).", False)], NUM, TXT)
    add_numbered_item(doc, 6, [("Applicable national legislation regarding waste prevention "
                                "and management (e.g. for transport, supervision, etc.) are "
                                "complied with. If veterinary certificates exist, these are to "
                                "be kept together with the commercial documents.", False)],
                      NUM, TXT)
    add_numbered_item(doc, 7, [("The supplied material is exclusively generated or occurred at "
                                "the signing Point of Origin.", False)], NUM, TXT)
    add_numbered_item(doc, 8, [("Auditors from Certification Bodies or from ISCC may, with or "
                                "without prior notice, verify on-site or by contacting the "
                                "Company (e.g. via telephone), whether the relevant ISCC PLUS "
                                "requirements are complied with and whether the statements "
                                "made in this self-declaration are correct. Auditors may be "
                                "accompanied by inspectors who monitor their activities.",
                                False)], NUM, TXT)
    add_numbered_item(doc, 9, [(ITEM_9, False)], NUM, TXT)
    add_numbered_item(doc, 10, [("This self-declaration and the information contained herein "
                                 "may be forwarded, including for review or further "
                                 "processing, by any relevant element of the supply chain, the "
                                 "Certification Body, ISCC, competent authorities or "
                                 "supervisory bodies, or, where legally required, any other "
                                 "institution or entity, to such bodies or to third parties "
                                 "acting on their behalf to ensure and enforce compliance.",
                                 False)], NUM, TXT)
    add_numbered_item(doc, 11, [(ITEM_ACK, False)], NUM, TXT)
    add_numbered_item(doc, 12, [(ITEM_WARRANT, False)], NUM, TXT)
    add_numbered_item(doc, 13, [(ITEM_PROVIDE, False)], NUM, TXT)
    add_numbered_item(doc, 14, [(ITEM_CORRECT, False)], NUM, TXT)
    add_numbered_item(doc, 15, [(ITEM_LAW, False)], NUM, TXT, after=18)

    # signature block (PDF: label x=40, field 99..233, signature ~250.., field ..546)
    add_signature_block(doc, [74, 130, 62, text_w - 74 - 130 - 62],
                        "{legalPerson}, {position}")

    # footnotes
    p = doc.add_paragraph()
    para_format(p, before=24, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "* " + FOOTNOTE_STAR, size=9)
    p = doc.add_paragraph()
    para_format(p, before=6, after=6)
    add_run(p, "1", size=9, superscript=True)
    add_run(p, " Example feedstocks: circular (including technical-circular): mixed plastic "
               "waste (MPW); bio-circular: Used Cooking Oil (UCO). Further explanation in "
               "ISCC PLUS material list", size=9)
    p = doc.add_paragraph()
    para_format(p, before=6, after=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_run(p, "This document neither replaces official delivery documents nor does it take "
               "precedence over national waste classification legislation. Although under "
               "ISCC PLUS the waste definition is based on Waste Framework Directive "
               "2008/98/EC (Article 3) and residue definition is based on Renewable Energy "
               "Directive II 2008/98/EU (Article 2), this document is not binding for the "
               "signature holder to comply with EU legislation if they are operating outside "
               "EU. Signature holder should only be aligned with the definition of waste and "
               "residues as provided above.", size=9)

    doc.save(path)
    declare_symbol_font(path)


# ================================================================== ISCC EU
def build_eu(path):
    # PDF: left 35.8pt, right 595.2-547.9=47.3pt, title top 71.5pt,
    # footer logo 780..815pt from top, copyright line at 819pt
    text_w = 512.1  # 547.9 - 35.8
    doc = new_document({"top": 25, "bottom": 27, "left": 12.6, "right": 16.7},
                       footer_mm=5.5)
    build_footer(doc, "Version 2.3, as of 30 September 2025", text_w, page_numbers=True)
    restart_page_numbers(doc)

    add_title_table(doc, "ISCC EU self-declaration",
                    "Point of Origin for Waste and Residues",
                    [224.4, text_w - 224.4])
    para_format(doc.add_paragraph(), after=0)

    # grid columns from PDF x: 35.8 | 286.6 | 506.4 | 547.9
    widths = [250.8, 219.8, 41.5]
    t = doc.add_table(rows=12, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(t, sz=4)
    set_table_cell_margins(t, top=40, bottom=40, left=100, right=100)
    set_table_fixed(t, widths)

    c = merge_row(t, 0, 0, 2)
    label_cell(c, INFO_HEADER, bold=True)
    set_row_height(t.rows[0], 14.4)

    simple_rows = [
        (1, "Site Name", "{storeName}"),
        (2, "Street address", "{address}"),
        (3, "City, Postcode", "{cityPostcode}"),
        (4, "Country", "{country}"),
    ]
    for r, label, tag in simple_rows:
        label_cell(t.cell(r, 0), label)
        placeholder_cell(merge_row(t, r, 1, 2), tag)
        set_row_height(t.rows[r], 16.3)

    label_cell(t.cell(5, 0), GEO_LABEL, extra_lines=GEO_HINT)
    placeholder_cell(merge_row(t, 5, 1, 2), "{geoCoordinates}")
    set_row_height(t.rows[5], 40)

    more_rows = [
        (6, "Phone number", "{phone}", 16.3),
        (7, "Maximum estimated capacity per year (in mt)", "{maxCapacity}", 16.6),
        (8, "Maximum estimated sustainable capacity per year (in mt)",
         "{maxSustainableCapacity}", 26),
    ]
    for r, label, tag, h in more_rows:
        label_cell(t.cell(r, 0), label)
        placeholder_cell(merge_row(t, r, 1, 2), tag)
        set_row_height(t.rows[r], h)

    # R9 material
    c = merge_row(t, 9, 0, 2)
    set_cell_valign(c, "top")
    p = c.paragraphs[0]
    para_format(p, after=4)
    add_run(p, MATERIAL_LABEL)
    p = c.add_paragraph()
    para_format(p, before=10, after=12, left=6)
    add_run(p, MATERIAL_TEXT)
    p = c.add_paragraph()
    para_format(p, before=2, after=0)
    add_run(p, MATERIAL_NOTE)
    set_row_height(t.rows[9], 70)

    # R10 volume statement + checkbox
    c = merge_row(t, 10, 0, 1)
    label_cell(c, "The amount of waste and residues generated by the Point of Origin is "
                  "five (5) or more metric tons per month")
    cb = t.cell(10, 2)
    set_cell_valign(cb, "center")
    p = cb.paragraphs[0]
    para_format(p, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "{minVolumeCheck}", font=CHECKBOX_FONT, east_asia=CHECKBOX_FONT)
    set_row_height(t.rows[10], 24)

    # R11 recipient
    label_cell(t.cell(11, 0), "Recipient of the waste or residue (Collecting Point)")
    placeholder_cell(merge_row(t, 11, 1, 2), "{collectionPoint}")
    set_row_height(t.rows[11], 34)

    for row in t.rows:
        set_row_cant_split(row)

    add_signing_sentence(doc)
    NUM, TXT = 4, 21  # PDF: number at x=42, text at x=55 (margin 35.8); "10." must fit
    add_numbered_item(doc, 1, [(ITEM_1, False)], NUM, TXT)
    add_numbered_item(doc, 2, [(ITEM_2, False)], NUM, TXT, after=3, keep_next=True)
    add_sub_paragraph(doc, [("A ", False), ("waste", True),
                            (" is any substance or object which the holder discards or "
                             "intends or is required to discard, excluding substances that "
                             "have been intentionally modified or contaminated in order to "
                             "meet this definition.", False)], TXT, after=3)
    add_sub_paragraph(doc, [("A ", False), ("residue", True),
                            (" is a substance that is not the end product that a production "
                             "process directly seeks to produce; it is not a primary aim of "
                             "the production process and the process has not been "
                             "deliberately modified to produce it.", False)], TXT)
    add_numbered_item(doc, 3, [("In case of residues that are directly generated by "
                                "agriculture, aquaculture, fisheries and forestry, the "
                                "material fulfils the land related sustainability requirements "
                                "laid down in Art. 29 of Directive (EU) 2023/2413 (REDIII) "
                                "amending Directive (EU) 2018/2001 (RED II).", False)],
                      NUM, TXT)
    add_numbered_item(doc, 4, [("The material supplied consists only of biomass defined as the "
                                "biodegradable fraction of products, waste and residues from "
                                "biological origin from agriculture (including vegetal and "
                                "animal substances), forestry and related industries including "
                                "fisheries and aquaculture, as well as the biodegradable "
                                "fraction of industrial and municipal waste.", False)],
                      NUM, TXT)
    add_numbered_item(doc, 5, [("Documentation of quantities supplied is available.", False)],
                      NUM, TXT)
    add_numbered_item(doc, 6, [("Applicable national legislations regarding waste prevention "
                                "and management (e.g. for transport, supervision, etc.) are "
                                "complied with. If veterinary certificates exist, these are to "
                                "be kept together with the commercial documents.", False)],
                      NUM, TXT)
    add_numbered_item(doc, 7, [("The supplied material is exclusively generated by the signing "
                                "Point of Origin.", False)], NUM, TXT)
    add_numbered_item(doc, 8, [("Auditors from Certification Bodies or from ISCC may, with or "
                                "without prior notice, verify on-site or by contacting the "
                                "Company (e.g. via telephone), whether the relevant ISCC EU "
                                "requirements are complied with and whether the statements "
                                "made in this self-declaration are correct. Auditors may be "
                                "accompanied by inspectors who monitor their activities.",
                                False)], NUM, TXT)
    add_numbered_item(doc, 9, [(ITEM_9, False)], NUM, TXT)
    add_numbered_item(doc, 10, [("This self-declaration or the information contained therein "
                                 "may be forwarded, including for review or further "
                                 "processing, by any relevant element of the supply chain, the "
                                 "Certification Body, ISCC or competent authorities or "
                                 "supervisory bodies, or, if legally required, by any other "
                                 "institution or entity, to each of the aforementioned bodies "
                                 "as well as to third parties who act on behalf of these "
                                 "bodies or entities to ensure and enforce compliance.",
                                 False)], NUM, TXT)
    add_numbered_item(doc, 11, [(ITEM_ACK, False)], NUM, TXT)
    add_numbered_item(doc, 12, [("The information contained in this self-declaration and the "
                                 "information referred to in Statement 11 may be forwarded to "
                                 "any data base operated by or on behalf of the European Union "
                                 "or any of its Member States, e.g. the Union Database for "
                                 "Biofuels (UDB), and to any Service Provider that provides "
                                 "access to or facilitates data handling in such database.",
                                 False)], NUM, TXT)
    add_numbered_item(doc, 13, [(ITEM_WARRANT, False)], NUM, TXT)
    add_numbered_item(doc, 14, [(ITEM_PROVIDE, False)], NUM, TXT)
    add_numbered_item(doc, 15, [(ITEM_CORRECT, False)], NUM, TXT)
    add_numbered_item(doc, 16, [(ITEM_LAW, False)], NUM, TXT, after=18)

    # signature block (PDF: Place,Date x=40, field 99..233, Signature x=278, field 330..558)
    add_signature_block(doc, [74, 130, 62, text_w - 74 - 130 - 62],
                        "{legalPerson}, {position}")

    p = doc.add_paragraph()
    para_format(p, before=24, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "*" + FOOTNOTE_STAR, size=8)

    doc.save(path)
    declare_symbol_font(path)


if __name__ == "__main__":
    os.makedirs(OUT_DIR, exist_ok=True)
    build_plus(os.path.join(OUT_DIR, "ISCC_PLUS.docx"))
    build_eu(os.path.join(OUT_DIR, "ISCC_EU.docx"))
    print("built", os.path.join(OUT_DIR, "ISCC_PLUS.docx"), os.path.join(OUT_DIR, "ISCC_EU.docx"))
